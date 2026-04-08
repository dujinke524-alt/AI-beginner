from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
# 注意：这里我们从 langchain_core 导入 Document 类，用于创建 LangChain 的文档对象
from langchain_core.documents import Document as LCDocument
from langchain_community.embeddings import DashScopeEmbeddings
from langchain_community.vectorstores import FAISS
import os

# --- 1. 读取 Word 文档 ---
# 为了避免混淆，我们将 Word 文档对象命名为 word_doc
word_doc = Document(r"sample.docx")

# --- 2. 提取并整合所有文本 ---
# 创建一个大列表，用来存放文档中所有的文本内容
all_text = []

# 提取所有段落文本
for para in word_doc.paragraphs:
    text = para.text.strip()
    if text:  # 只添加非空段落
        all_text.append(text)

# 提取所有表格文本
for table in word_doc.tables:
    for row in table.rows:
        # 将一行中的所有单元格文本合并成一个字符串
        row_text = " ".join([cell.text.strip() for cell in row.cells])
        if row_text:  # 只添加非空行
            all_text.append(row_text)

# 将所有文本用换行符连接成一个大的字符串
full_text = "\n".join(all_text)

# --- 3. 使用 TextSplitter 分割文本 ---
splitter = RecursiveCharacterTextSplitter(
    chunk_size=100,
    chunk_overlap=20
)
# 对整合后的完整文本进行分割
text_chunks = splitter.split_text(full_text)

# 打印分割后的文本块，检查效果
print("--- 分割后的文本块示例 ---")
for i, chunk in enumerate(text_chunks[:2], 1):  # 只打印前2块看看
    print(f"--- 第{i}块 ---")
    print(chunk)
    print("\n")

# --- 4. 将文本块转换为 LangChain 的 Document 对象 ---
# from_documents 方法需要一个 Document 对象的列表
# 我们遍历 text_chunks，为每个文本块创建一个 LCDocument 对象
docs_for_vectorstore = [LCDocument(page_content=chunk) for chunk in text_chunks]

# --- 5. 使用千问 Embedding ---
embeddings = DashScopeEmbeddings(
    model="text-embedding-v4",
    # 确保你已经在系统环境变量中设置了 DASHSCOPE_API_KEY
    # 如果没有，可以直接在这里写死：dashscope_api_key="sk-你的key"
    dashscope_api_key="sk-你的key"
)

# --- 6. 放进 FAISS 向量库 ---
# 这里传入的是我们刚刚创建的 Document 对象列表
vectorstore = FAISS.from_documents(docs_for_vectorstore, embeddings)

# --- 7. 提问题并搜索 ---
print('输入你的问题:')
question =  input()

# 相似度搜索，返回最相关的 3 段
results = vectorstore.similarity_search(question, k=3)

# --- 8. 打印结果 ---
print("\n\n--- 搜索结果 ---")
# 为了避免混淆，我们将搜索结果中的对象命名为 result_doc
for i, result_doc in enumerate(results, start=1):
    print(f"--- 第{i}段相关原文 ---")
    print(result_doc)
    print("\n")
